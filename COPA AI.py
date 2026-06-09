
import os
import re
import io
import json
import time
import uuid
import hashlib
import threading
import contextlib
import unicodedata
from datetime import datetime

import pandas as pd
import tkinter as tk
from tkinter import filedialog, scrolledtext, ttk, messagebox
import tkinter.font as tkfont

try:
    from PIL import Image
except Exception:
    Image = None

try:
    import plotly
except Exception:
    plotly = None

from docx import Document

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg

import ollama

APP_TITLE = 'COPA AI - Python Code Interpreter'
APP_GEOMETRY = '1380x820'
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
MEMORY_FILE = os.path.join(BASE_DIR, 'financial_assistant_memory.json')
DEFAULT_MODEL = 'gemma4:e2b'
MAX_CHAT_HISTORY = 300
MAX_SAMPLE_ROWS = 15
MAX_CONVERSATION_MESSAGES = 300
MAX_RECENT_CONVERSATIONS = 80000
INPUT_FONT = ('Segoe UI', 12)
INPUT_MIN_LINES = 3
INPUT_MAX_LINES = 8
INPUT_MIN_HEIGHT = 96
INPUT_MAX_HEIGHT = 260
ATTACH_BAR_HEIGHT = 34
DEBUG_PANEL_WIDTH = 360
SUPPORTED_TEXT_EXTS = {'.txt'}
SUPPORTED_WORD_EXTS = {'.docx'}
SUPPORTED_CSV_EXTS = {'.csv'}
BLOCKED_EXTS = {'.pdf', '.png', '.jpg', '.jpeg', '.webp', '.bmp'}
SUPPORTED_EXCEL_EXTS = {'.xlsx', '.xls'}

COLORS = {
    'bg': '#0f1117', 'bg_soft': '#121622', 'sidebar': '#1b1c1f',
    'sidebar_hover': '#25262b', 'sidebar_selected': '#2b2d34',
    'card': '#1f2430', 'card_light': '#242a36', 'text': '#eef2f7',
    'muted': '#a7b0be', 'muted_2': '#7e8795', 'accent': '#8ab4f8',
    'accent_2': '#77e0a6', 'danger': '#ff6b6b', 'warning': '#f2b84b',
    'input': '#202124', 'input_border': '#3b4252', 'chip': '#2a3140',
    'white': '#ffffff',
}


def build_reader_code(filepath, max_chars=8000):
    ext = os.path.splitext(filepath)[1].lower()
    safe_path = filepath.replace('\\', '\\\\')
    if ext in ['.xlsx', '.xls']:
        return f'''
import pandas as pd
filepath = r"{safe_path}"
excel_file = pd.ExcelFile(filepath)
print("Archivo Excel detectado.")
print("Hojas disponibles:", excel_file.sheet_names)
sheet = excel_file.sheet_names[0]
df = pd.read_excel(filepath, sheet_name=sheet).fillna("")
print("Hoja leída:", sheet)
print("Filas:", len(df))
print("Columnas:", list(df.columns))
print("\\nMuestra de datos:")
print(df.head(10).to_string(index=False))
'''
    if ext == '.csv':
        return f'''
import pandas as pd
filepath = r"{safe_path}"
df = pd.read_csv(filepath).fillna("")
print("Archivo CSV detectado.")
print("Filas:", len(df))
print("Columnas:", list(df.columns))
print("\\nMuestra de datos:")
print(df.head(10).to_string(index=False))
'''
    if ext == '.txt':
        return f'''
filepath = r"{safe_path}"
with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
    content = f.read()
print("Archivo TXT detectado.")
print("Caracteres extraídos:", len(content))
print("\\nContenido extraído:")
print(content[:{max_chars}])
'''
    if ext == '.docx':
        return f'''
from docx import Document
filepath = r"{safe_path}"
doc = Document(filepath)
parts = []
for p in doc.paragraphs:
    text = p.text.strip()
    if text:
        parts.append(text)
for table in doc.tables:
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if any(cells):
            parts.append(" | ".join(cells))
content = "\\n".join(parts)
print("Archivo Word DOCX detectado.")
print("Caracteres extraídos:", len(content))
print("\\nContenido extraído:")
print(content[:{max_chars}])
'''
    return 'print("Tipo de archivo no soportado para extracción automática.")\n'


def normalize_text(value):
    if value is None:
        return ''
    value = str(value).strip().lower()
    value = unicodedata.normalize('NFKD', value)
    return ''.join(ch for ch in value if not unicodedata.combining(ch))


def safe_now_iso():
    return datetime.now().isoformat(timespec='seconds')


def make_signature(filepath, sheet_name, columns):
    raw = f'{filepath}|{sheet_name}|{"|".join(map(str, columns))}'
    return hashlib.sha256(raw.encode('utf-8')).hexdigest()


def first_nonempty_line(text, max_len=52):
    clean = re.sub(r'\s+', ' ', (text or '').strip())
    if not clean:
        return 'Nueva conversación'
    return clean if len(clean) <= max_len else clean[: max_len - 1].rstrip() + '…'


def get_chunk_message_dict(chunk):
    try:
        if isinstance(chunk, dict):
            return chunk.get('message', {}) or {}
    except Exception:
        pass
    try:
        msg = getattr(chunk, 'message', None)
        if msg is None:
            return {}
        if isinstance(msg, dict):
            return msg
        return {
            'role': getattr(msg, 'role', None),
            'content': getattr(msg, 'content', '') or '',
            'thinking': getattr(msg, 'thinking', '') or '',
        }
    except Exception:
        return {}


def classify_attachment(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    if ext in SUPPORTED_EXCEL_EXTS:
        return 'excel'
    if ext in SUPPORTED_CSV_EXTS:
        return 'csv'
    if ext in SUPPORTED_TEXT_EXTS:
        return 'text'
    if ext in SUPPORTED_WORD_EXTS:
        return 'word'
    if ext in BLOCKED_EXTS:
        return 'blocked'
    return 'unknown'


class MemoryManager:
    def __init__(self, filepath=MEMORY_FILE):
        self.filepath = filepath
        self.data = self._load()

    def _default_data(self):
        return {
            'schema_version': 2,
            'datasets': {},
            'global_columns': {},
            'global': {'default_model': DEFAULT_MODEL, 'active_conversation_id': None, 'memory_notes': []},
            'conversations': {},
        }

    def _load(self):
        data = self._default_data()
        if os.path.exists(self.filepath):
            try:
                with open(self.filepath, 'r', encoding='utf-8') as f:
                    loaded = json.load(f)
                if isinstance(loaded, dict):
                    data.update(loaded)
            except Exception as e:
                print(f'Aviso al cargar memoria: {e}')
        data.setdefault('schema_version', 2)
        data.setdefault('datasets', {})
        data.setdefault('global_columns', {})
        data.setdefault('global', {})
        data['global'].setdefault('default_model', DEFAULT_MODEL)
        data['global'].setdefault('active_conversation_id', None)
        data['global'].setdefault('memory_notes', [])
        data.setdefault('conversations', {})
        return data

    def save(self):
        tmp_path = f'{self.filepath}.tmp'
        try:
            os.makedirs(os.path.dirname(self.filepath), exist_ok=True)
            with open(tmp_path, 'w', encoding='utf-8') as f:
                json.dump(self.data, f, ensure_ascii=False, indent=2)
            os.replace(tmp_path, self.filepath)
        except Exception as e:
            print(f'Error al guardar memoria: {e}')
            try:
                if os.path.exists(tmp_path):
                    os.remove(tmp_path)
            except Exception:
                pass

    def create_conversation(self, title='Nueva conversación', save_now=True):
        conv_id = uuid.uuid4().hex
        conv = {
            'id': conv_id, 'title': title or 'Nueva conversación',
            'created_at': safe_now_iso(), 'updated_at': safe_now_iso(),
            'messages': [], 'pinned': False, 'summary': '', 'dataset_signatures': [],
        }
        self.data.setdefault('conversations', {})[conv_id] = conv
        self.data.setdefault('global', {})['active_conversation_id'] = conv_id
        if save_now:
            self.save()
        return conv

    def get_active_conversation_id(self):
        return self.data.get('global', {}).get('active_conversation_id')

    def set_active_conversation(self, conv_id):
        if conv_id in self.data.get('conversations', {}):
            self.data.setdefault('global', {})['active_conversation_id'] = conv_id
            self.save()
            return True
        return False

    def get_conversation(self, conv_id=None):
        convs = self.data.setdefault('conversations', {})
        conv_id = conv_id or self.get_active_conversation_id()
        if conv_id and conv_id in convs:
            return convs[conv_id]
        return self.create_conversation(save_now=True)

    def list_conversations(self, query=''):
        convs = list(self.data.get('conversations', {}).values())
        q = normalize_text(query)
        if q:
            convs = [c for c in convs if q in normalize_text(c.get('title', '')) or q in normalize_text(c.get('summary', ''))]
        convs.sort(key=lambda c: (bool(c.get('pinned')), c.get('updated_at', '')), reverse=True)
        return convs[:MAX_RECENT_CONVERSATIONS]

    def append_message(self, role, content, conv_id=None, extra=None):
        conv = self.get_conversation(conv_id)
        msg = {'role': role, 'content': content, 'created_at': safe_now_iso()}
        if extra:
            msg.update(extra)
        conv.setdefault('messages', []).append(msg)
        conv['messages'] = conv['messages'][-MAX_CONVERSATION_MESSAGES:]
        conv['updated_at'] = safe_now_iso()
        if role == 'user' and (not conv.get('title') or conv.get('title') == 'Nueva conversación'):
            conv['title'] = first_nonempty_line(content)
        self.save()
        return conv

    def delete_conversation(self, conv_id):
        convs = self.data.setdefault('conversations', {})
        if conv_id in convs:
            del convs[conv_id]
            if self.get_active_conversation_id() == conv_id:
                self.data['global']['active_conversation_id'] = None
            self.save()
            return True
        return False

    def remember_dataset_for_active_chat(self, signature):
        conv = self.get_conversation()
        signatures = conv.setdefault('dataset_signatures', [])
        if signature and signature not in signatures:
            signatures.append(signature)
            conv['updated_at'] = safe_now_iso()
            self.save()

    def get_memory_notes(self):
        return [n for n in self.data.get('global', {}).get('memory_notes', []) if isinstance(n, dict)]

    def add_memory_note(self, text, source='manual'):
        text = (text or '').strip()
        if not text:
            return
        notes = self.data.setdefault('global', {}).setdefault('memory_notes', [])
        notes.append({'text': text[:500], 'source': source, 'created_at': safe_now_iso()})
        self.data['global']['memory_notes'] = notes[-80:]
        self.save()

    def get_dataset(self, signature):
        return self.data.get('datasets', {}).get(signature, {})

    def update_dataset(self, signature, payload):
        self.data.setdefault('datasets', {})[signature] = payload
        self.save()

    def get_global_column_meta(self, col_name):
        return self.data.get('global_columns', {}).get(col_name)

    def update_global_dictionary(self, definitions):
        self.data.setdefault('global_columns', {})
        for col, meta in (definitions or {}).items():
            self.data['global_columns'][col] = meta
        self.save()


class ExcelProcessor:
    def __init__(self, filepath, sheet_name, column_definitions=None):
        self.filepath = filepath
        self.filename = os.path.basename(filepath)
        self.sheet_name = sheet_name
        self.df = None
        self.signature = None
        self.column_definitions = column_definitions or {}
        self.column_roles = {}
        self.column_meanings = {}
        self.precomputed = {}

    def load(self):
        ext = os.path.splitext(self.filepath)[1].lower()
        engine = 'openpyxl' if ext == '.xlsx' else 'xlrd' if ext == '.xls' else None
        self.df = pd.read_excel(self.filepath, sheet_name=self.sheet_name, engine=engine).fillna('')
        self.df.columns = [str(c).strip() for c in self.df.columns]
        self.signature = make_signature(self.filepath, self.sheet_name, self.df.columns.tolist())
        for col in self.df.columns:
            if col in self.column_definitions:
                self.column_roles[col] = self.column_definitions[col].get('role', '')
                self.column_meanings[col] = self.column_definitions[col].get('meaning', '')
            else:
                self.column_roles[col] = ''
                self.column_meanings[col] = ''
        self._infer_roles()
        self._compute_predefined_metrics()

    def _infer_roles(self):
        for col in self.df.columns:
            if self.column_roles.get(col):
                continue
            series = self.df[col]
            if pd.api.types.is_datetime64_any_dtype(series):
                self.column_roles[col] = 'fecha'
                continue
            num_series = pd.to_numeric(series, errors='coerce')
            if len(series) > 0 and num_series.notna().sum() > len(series) * 0.4:
                self.column_roles[col] = 'valores'
                continue
            unique_count = series.nunique()
            if len(series) > 0 and (unique_count <= 25 or unique_count < len(series) * 0.15):
                self.column_roles[col] = 'categoría principal'
            else:
                self.column_roles[col] = 'texto descriptivo'

    def apply_user_definitions(self, definitions):
        if not definitions:
            return
        for col, meta in definitions.items():
            if col in self.df.columns:
                self.column_roles[col] = meta.get('role', 'unknown')
                self.column_meanings[col] = meta.get('meaning', '')
        self._compute_predefined_metrics()

    def _compute_predefined_metrics(self):
        self.precomputed = {
            'rows': len(self.df),
            'columns_count': len(self.df.columns),
            'sample_rows': self.df.head(MAX_SAMPLE_ROWS).to_string(index=False),
        }

    def build_llm_context(self):
        lines = [f'Archivo: {self.filename} (Hoja: {self.sheet_name})', 'DICCIONARIO DE DATOS (Metadata):']
        for col in self.df.columns:
            role = self.column_roles.get(col, 'unknown')
            meaning = self.column_meanings.get(col, '')
            lines.append(f'- {col}: Tipo={role} | Significado={meaning}')
        lines.append(f"\nMuestra de datos ({MAX_SAMPLE_ROWS} filas):\n{self.precomputed['sample_rows']}")
        return '\n'.join(lines)

    def build_memory_payload(self):
        return {
            'file': self.filename, 'sheet': self.sheet_name, 'signature': self.signature,
            'column_definitions': {col: {'role': self.column_roles.get(col, ''), 'meaning': self.column_meanings.get(col, '')} for col in self.df.columns},
            'updated_at': safe_now_iso(),
        }


class SheetSelectorDialog(tk.Toplevel):
    def __init__(self, parent, sheets):
        super().__init__(parent)
        self.title('Seleccionar hoja')
        self.geometry('420x340')
        self.configure(bg=COLORS['bg_soft'])
        self.resizable(False, False)
        self.selected_sheet = None
        self.transient(parent)
        self.grab_set()
        tk.Label(self, text='El archivo tiene varias hojas.\nSelecciona la hoja a analizar:', font=('Segoe UI', 11, 'bold'), fg=COLORS['text'], bg=COLORS['bg_soft']).pack(pady=14)
        self.listbox = tk.Listbox(self, font=('Consolas', 11), bg=COLORS['card'], fg=COLORS['text'], selectbackground=COLORS['accent'], selectforeground='#0b1020', highlightthickness=0, relief=tk.FLAT)
        self.listbox.pack(expand=True, fill=tk.BOTH, padx=16, pady=8)
        for s in sheets:
            self.listbox.insert(tk.END, s)
        if sheets:
            self.listbox.selection_set(0)
        btn_frame = tk.Frame(self, bg=COLORS['bg_soft'])
        btn_frame.pack(pady=14)
        tk.Button(btn_frame, text='Aceptar', command=self.on_ok, bg=COLORS['accent'], fg='#08111f', width=12, relief=tk.FLAT, font=('Segoe UI', 10, 'bold')).pack(side=tk.LEFT, padx=8)
        tk.Button(btn_frame, text='Cancelar', command=self.on_cancel, bg=COLORS['card_light'], fg=COLORS['text'], width=12, relief=tk.FLAT, font=('Segoe UI', 10, 'bold')).pack(side=tk.LEFT, padx=8)
        self.protocol('WM_DELETE_WINDOW', self.on_cancel)

    def on_ok(self):
        selection = self.listbox.curselection()
        if selection:
            self.selected_sheet = self.listbox.get(selection[0])
        self.destroy()

    def on_cancel(self):
        self.selected_sheet = None
        self.destroy()


class ColumnDefinitionDialog(tk.Toplevel):
    def __init__(self, parent, processor):
        super().__init__(parent)
        self.title('Diccionario de Datos')
        self.geometry('1020x640')
        self.configure(bg=COLORS['bg_soft'])
        self.processor = processor
        self.result = {}
        self.transient(parent)
        self.grab_set()
        tk.Label(self, text=f'Configurando: {processor.filename}\nValida o completa el significado de las columnas para mejorar la precisión de la IA.', bg=COLORS['bg_soft'], fg=COLORS['text'], font=('Segoe UI', 11, 'bold'), justify=tk.LEFT).pack(anchor='w', padx=16, pady=14)
        container = tk.Frame(self, bg=COLORS['bg_soft'])
        container.pack(expand=True, fill=tk.BOTH, padx=14, pady=8)
        canvas = tk.Canvas(container, bg=COLORS['bg_soft'], highlightthickness=0)
        scrollbar = ttk.Scrollbar(container, orient='vertical', command=canvas.yview)
        self.scrollable_frame = tk.Frame(canvas, bg=COLORS['bg_soft'])
        self.scrollable_frame.bind('<Configure>', lambda e: canvas.configure(scrollregion=canvas.bbox('all')))
        canvas.create_window((0, 0), window=self.scrollable_frame, anchor='nw')
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side=tk.LEFT, expand=True, fill=tk.BOTH)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        hdr = tk.Frame(self.scrollable_frame, bg=COLORS['card_light'])
        hdr.pack(fill=tk.X, pady=(0, 6))
        for i, text in enumerate(['Columna', 'Tipo detectado', 'Diccionario de datos / Significado']):
            tk.Label(hdr, text=text, width=[28, 28, 48][i], anchor='w', bg=COLORS['card_light'], fg=COLORS['text'], font=('Segoe UI', 10, 'bold')).grid(row=0, column=i, padx=6, pady=6)
        self.widgets = {}
        for col in self.processor.df.columns:
            row = tk.Frame(self.scrollable_frame, bg=COLORS['bg_soft'])
            row.pack(fill=tk.X, pady=3)
            tk.Label(row, text=col, width=28, anchor='w', bg=COLORS['bg_soft'], fg=COLORS['text'], font=('Consolas', 10)).grid(row=0, column=0, padx=6)
            inferred_role = self.processor.column_roles.get(col, 'desconocido')
            tk.Label(row, text=inferred_role.upper(), width=28, anchor='w', bg=COLORS['bg_soft'], fg=COLORS['accent'], font=('Consolas', 9, 'bold')).grid(row=0, column=1, padx=6)
            meaning_var = tk.StringVar(value=self.processor.column_meanings.get(col, ''))
            ent = tk.Entry(row, textvariable=meaning_var, width=62, font=('Consolas', 10), bg=COLORS['input'], fg=COLORS['text'], insertbackground=COLORS['text'], relief=tk.FLAT)
            ent.grid(row=0, column=2, sticky='we', padx=6, ipady=4)
            self.widgets[col] = {'role': inferred_role, 'meaning_var': meaning_var}
        btn_frame = tk.Frame(self, bg=COLORS['bg_soft'])
        btn_frame.pack(fill=tk.X, padx=16, pady=14)
        tk.Button(btn_frame, text='Guardar diccionario', command=self.on_ok, bg=COLORS['accent_2'], fg='#08111f', font=('Segoe UI', 10, 'bold'), relief=tk.FLAT, padx=16, pady=7).pack(side=tk.RIGHT, padx=5)

    def on_ok(self):
        for col, obj in self.widgets.items():
            self.result[col] = {'role': obj['role'], 'meaning': obj['meaning_var'].get().strip()}
        self.destroy()


class FinancialAssistantApp:
    def __init__(self, root):
        self.root = root
        self.root.title(APP_TITLE)
        self.root.geometry(APP_GEOMETRY)
        self.root.configure(bg=COLORS['bg'])
        self._fullscreen_start()
        self.memory = MemoryManager()
        self.active_conversation = self.memory.get_conversation()
        self.history = self._messages_for_llm()
        self.dfs = []
        self.processors = []
        self.pending_attachments = []
        self.last_sent_attachments = []
        self.current_filepath = None
        self.live_answer_started = False
        self.last_chunk = ''
        self.cancel_event = threading.Event()
        self.is_busy = False
        self.current_run_id = 0
        self.thinking_panel_visible = True
        self.system_prompt_text = '''
Eres un Científico de Datos y Analista Financiero. Tienes acceso directo a una lista de DataFrames cargados en la variable `dfs`.
El archivo más reciente y principal siempre está en `dfs[-1]`, y por conveniencia también está mapeado en la variable `df`.

REGLAS CRÍTICAS:
1. REVISA SIEMPRE el diccionario de datos antes de operar.
2. Antes de operar matemáticamente verifica tipos de datos (`df.dtypes`). Si una columna numérica está como string, conviértela con `pd.to_numeric(..., errors='coerce')`.
3. Incluye `print()` con resultados intermedios relevantes para auditoría.
4. Para preguntas analíticas con datos cargados, genera un bloque de código Python usando pandas, matplotlib o plotly cuando aplique.
5. El código debe estar estrictamente dentro de un bloque ```python y ```.
6. Si recibes salida de código ejecutado, explica el resultado en lenguaje ejecutivo sin volver a escribir código.
7. Si el archivo es PDF o imagen, indica que este entorno no procesa ese formato y recomienda Microsoft Copilot para ese análisis.
8. Si el usuario adjunta archivos, COPA AI debe analizarlos mediante código Python local. Nunca asumas datos no extraídos o no cargados.

REGLAS OBLIGATORIAS PARA CÓDIGO PYTHON:
1. Siempre debes generar código Python ejecutable.
2. Si hay un archivo adjunto, NO inventes rutas ni nombres de archivo. Usa exactamente la variable `filepath` proporcionada por COPA AI.
3. Si `filepath` está disponible, úsalo para leer el archivo localmente. Si `filepath` es None pero existe `df`, trabaja con `df`.
4. Siempre imprime resultados importantes con `print()`.
5. Matplotlib: usa `fig, ax = plt.subplots(...)`, no uses `plt.close()`, no dependas de `plt.show()`, deja figuras abiertas.
6. Plotly: guarda cada figura como HTML local con `fig.write_html(...)`, ábrela con `webbrowser.open("file://" + html_path)`, imprime la ruta.
7. Al finalizar una ejecución exitosa, imprime: `Ejecución completada correctamente.`
'''
        self.setup_ui()
        self.load_active_conversation_into_chat()
        self.root.protocol('WM_DELETE_WINDOW', self.on_close)

    def _fullscreen_start(self):
        try:
            self.root.state('zoomed')
        except Exception:
            try:
                self.root.attributes('-zoomed', True)
            except Exception:
                self.root.attributes('-fullscreen', True)

    def _run_reader_sandbox(self, code_string):
        import traceback
        output_buffer = io.StringIO()
        error_buffer = io.StringIO()
        env = {'pd': pd, 'os': os, 'Document': Document}
        try:
            with contextlib.redirect_stdout(output_buffer), contextlib.redirect_stderr(error_buffer):
                exec(code_string, env)
            output = output_buffer.getvalue()
            stderr = error_buffer.getvalue()
            if stderr.strip():
                output += '\n\nSTDERR:\n' + stderr
            return True, output
        except Exception:
            return False, 'Error ejecutando extractor:\n' + traceback.format_exc()

    def setup_ui(self):
        self.main = tk.Frame(self.root, bg=COLORS['bg'])
        self.main.pack(expand=True, fill=tk.BOTH)
        self.sidebar = tk.Frame(self.main, bg=COLORS['sidebar'], width=300)
        self.sidebar.pack(side=tk.LEFT, fill=tk.Y)
        self.sidebar.pack_propagate(False)
        self.content = tk.Frame(self.main, bg=COLORS['bg'])
        self.content.pack(side=tk.LEFT, expand=True, fill=tk.BOTH)
        self.content.grid_rowconfigure(1, weight=1)
        self.content.grid_columnconfigure(0, weight=1)
        self._setup_sidebar()
        self._setup_chat_area()
        self.refresh_recent_list()

    def _setup_sidebar(self):
        brand = tk.Frame(self.sidebar, bg=COLORS['sidebar'])
        brand.pack(fill=tk.X, padx=12, pady=(12, 8))
        tk.Label(brand, text='COPA AI', fg=COLORS['text'], bg=COLORS['sidebar'], font=('Segoe UI', 13, 'bold')).pack(side=tk.LEFT)
        tk.Label(brand, text='local', fg=COLORS['muted'], bg=COLORS['sidebar'], font=('Segoe UI', 9)).pack(side=tk.RIGHT)
        tk.Button(self.sidebar, text='＋  Nueva conversación', command=self.new_conversation, bg=COLORS['sidebar_hover'], fg=COLORS['text'], activebackground=COLORS['sidebar_selected'], activeforeground=COLORS['text'], relief=tk.FLAT, anchor='w', padx=14, font=('Segoe UI', 10, 'bold'), height=2).pack(fill=tk.X, padx=10, pady=(4, 8))
        search_frame = tk.Frame(self.sidebar, bg=COLORS['sidebar'])
        search_frame.pack(fill=tk.X, padx=10, pady=(0, 12))
        tk.Label(search_frame, text='⌕', bg=COLORS['sidebar'], fg=COLORS['muted'], font=('Segoe UI', 12)).pack(side=tk.LEFT)
        self.search_var = tk.StringVar()
        self.search_var.trace_add('write', lambda *_: self.refresh_recent_list())
        tk.Entry(search_frame, textvariable=self.search_var, bg=COLORS['sidebar'], fg=COLORS['text'], insertbackground=COLORS['text'], relief=tk.FLAT, font=('Segoe UI', 10)).pack(side=tk.LEFT, expand=True, fill=tk.X, padx=(6, 0), ipady=5)
        tk.Label(self.sidebar, text='Recientes', fg=COLORS['muted'], bg=COLORS['sidebar'], font=('Segoe UI', 9, 'bold'), anchor='w').pack(fill=tk.X, padx=12, pady=(6, 4))
        self.recent_canvas = tk.Canvas(self.sidebar, bg=COLORS['sidebar'], highlightthickness=0)
        self.recent_scroll = ttk.Scrollbar(self.sidebar, orient='vertical', command=self.recent_canvas.yview)
        self.recent_frame = tk.Frame(self.recent_canvas, bg=COLORS['sidebar'])
        self.recent_frame.bind('<Configure>', lambda e: self.recent_canvas.configure(scrollregion=self.recent_canvas.bbox('all')))
        self.recent_canvas.create_window((0, 0), window=self.recent_frame, anchor='nw')
        self.recent_canvas.configure(yscrollcommand=self.recent_scroll.set)
        self.recent_canvas.pack(side=tk.LEFT, expand=True, fill=tk.BOTH, padx=(8, 0), pady=(0, 8))
        self.recent_scroll.pack(side=tk.RIGHT, fill=tk.Y, pady=(0, 8))
        footer = tk.Frame(self.sidebar, bg=COLORS['sidebar'])
        footer.pack(fill=tk.X, padx=10, pady=10)
        tk.Button(footer, text='📥 Cargar Excel', command=self.load_excel, bg=COLORS['accent_2'], fg='#08111f', relief=tk.FLAT, font=('Segoe UI', 10, 'bold'), height=2).pack(fill=tk.X, pady=(0, 8))
        tk.Button(footer, text='🧠 Recordar nota', command=self.remember_current_note, bg=COLORS['card_light'], fg=COLORS['text'], relief=tk.FLAT, font=('Segoe UI', 10, 'bold'), height=2).pack(fill=tk.X)

    def _setup_chat_area(self):
        top = tk.Frame(self.content, bg=COLORS['bg'], height=64)
        top.grid(row=0, column=0, sticky='ew', padx=24, pady=(16, 8))
        top.grid_propagate(False)
        top.grid_columnconfigure(0, weight=1)
        self.lbl_chat_title = tk.Label(top, text='Nueva conversación', fg=COLORS['text'], bg=COLORS['bg'], font=('Segoe UI', 16, 'bold'), anchor='w')
        self.lbl_chat_title.grid(row=0, column=0, sticky='w')
        self.show_thinking_var = tk.BooleanVar(value=True)
        tk.Button(top, text='Limpiar debug', command=self.clear_thinking_panel, bg=COLORS['card_light'], fg=COLORS['text'], relief=tk.FLAT, font=('Segoe UI', 10, 'bold'), padx=10, pady=6).grid(row=0, column=1, sticky='e', padx=(8, 0))
        tk.Checkbutton(top, text='Thinking', variable=self.show_thinking_var, command=self.toggle_thinking_panel, bg=COLORS['bg'], fg=COLORS['muted'], selectcolor=COLORS['bg'], activebackground=COLORS['bg'], activeforeground=COLORS['text'], font=('Segoe UI', 10)).grid(row=0, column=2, sticky='e', padx=(8, 0))
        self.body = tk.Frame(self.content, bg=COLORS['bg'])
        self.body.grid(row=1, column=0, sticky='nsew', padx=24, pady=(0, 14))
        self.body.grid_rowconfigure(0, weight=1)
        self.body.grid_columnconfigure(0, weight=1)
        self.body.grid_columnconfigure(1, minsize=DEBUG_PANEL_WIDTH)
        self.chat_frame = tk.Frame(self.body, bg=COLORS['bg'])
        self.chat_frame.grid(row=0, column=0, sticky='nsew', padx=(0, 14))
        self.chat_frame.grid_rowconfigure(0, weight=1)
        self.chat_frame.grid_rowconfigure(1, minsize=INPUT_MIN_HEIGHT)
        self.chat_frame.grid_columnconfigure(0, weight=1)
        self.chat_display = scrolledtext.ScrolledText(self.chat_frame, wrap=tk.WORD, bg=COLORS['bg'], fg=COLORS['text'], insertbackground=COLORS['text'], font=('Segoe UI', 11), state=tk.DISABLED, relief=tk.FLAT, padx=14, pady=14)
        self.chat_display.grid(row=0, column=0, sticky='nsew')
        self.input_outer = tk.Frame(self.chat_frame, bg=COLORS['input_border'], height=INPUT_MIN_HEIGHT)
        self.input_outer.grid(row=1, column=0, sticky='ew', pady=(12, 0))
        self.input_outer.grid_propagate(False)
        self.input_outer.grid_columnconfigure(1, weight=1)
        self.input_outer.grid_rowconfigure(0, weight=1)
        self.btn_attach = tk.Button(self.input_outer, text='＋', command=self.attach_files, bg=COLORS['card_light'], fg=COLORS['text'], activebackground=COLORS['sidebar_selected'], activeforeground=COLORS['text'], font=('Segoe UI', 14, 'bold'), width=3, relief=tk.FLAT)
        self.btn_attach.grid(row=0, column=0, sticky='ns', padx=(1, 8), pady=1)
        self.input_inner = tk.Frame(self.input_outer, bg=COLORS['input'], height=INPUT_MIN_HEIGHT - 2)
        self.input_inner.grid(row=0, column=1, sticky='nsew', padx=0, pady=1)
        self.input_inner.grid_propagate(False)
        self.input_inner.grid_columnconfigure(0, weight=1)
        self.input_inner.grid_rowconfigure(1, weight=1)
        self.attachments_bar = tk.Frame(self.input_inner, bg=COLORS['input'], height=0)
        self.attachments_bar.grid(row=0, column=0, sticky='ew', padx=8, pady=(6, 0))
        self.attachments_bar.grid_remove()
        self.input_box = tk.Text(self.input_inner, height=INPUT_MIN_LINES, font=INPUT_FONT, bg=COLORS['input'], fg=COLORS['text'], insertbackground=COLORS['text'], relief=tk.FLAT, borderwidth=0, highlightthickness=0, padx=14, pady=12, wrap=tk.WORD, undo=True, autoseparators=True)
        self.input_box.grid(row=1, column=0, sticky='nsew')
        self.input_box.bind('<Return>', self.handle_return)
        self.input_box.bind('<KeyRelease>', self._auto_resize_input_box)
        self.input_box.bind('<Configure>', self._auto_resize_input_box)
        self.btn_send = tk.Button(self.input_outer, text='Enviar', command=self.send_message, bg=COLORS['accent'], fg='#08111f', activebackground=COLORS['accent'], activeforeground='#08111f', font=('Segoe UI', 11, 'bold'), width=12, relief=tk.FLAT)
        self.btn_send.grid(row=0, column=2, sticky='ns', padx=(8, 1), pady=1)
        self.debug_frame = tk.Frame(self.body, bg=COLORS['card'], width=DEBUG_PANEL_WIDTH)
        self.debug_frame.grid(row=0, column=1, sticky='nsew')
        self.debug_frame.grid_propagate(False)
        self.debug_frame.grid_rowconfigure(1, weight=1)
        self.debug_frame.grid_columnconfigure(0, weight=1)
        debug_header = tk.Frame(self.debug_frame, bg=COLORS['card_light'])
        debug_header.grid(row=0, column=0, sticky='ew')
        tk.Label(debug_header, text='Thinking / Debug', fg=COLORS['text'], bg=COLORS['card_light'], font=('Segoe UI', 11, 'bold')).pack(side=tk.LEFT, padx=12, pady=10)
        self.lbl_thinking_status = tk.Label(debug_header, text='Inactivo', fg=COLORS['muted'], bg=COLORS['card_light'], font=('Segoe UI', 10))
        self.lbl_thinking_status.bind('<Button-1>', lambda e: self.cancel_analysis())
        self.lbl_thinking_status.pack(side=tk.RIGHT, padx=12)
        self.thinking_display = scrolledtext.ScrolledText(self.debug_frame, wrap=tk.WORD, bg=COLORS['card'], fg=COLORS['muted'], font=('Consolas', 10), state=tk.DISABLED, relief=tk.FLAT, padx=10, pady=10)
        self.thinking_display.grid(row=1, column=0, sticky='nsew')

    def attach_files(self):
        filetypes = [('Archivos soportados', '*.xlsx *.xls *.csv *.txt *.docx'), ('Excel', '*.xlsx *.xls'), ('CSV', '*.csv'), ('Texto', '*.txt'), ('Word', '*.docx'), ('Todos', '*.*')]
        paths = filedialog.askopenfilenames(title='Adjuntar archivos', filetypes=filetypes)
        for path in paths:
            self._add_attachment(path)
        self._refresh_attachments_bar()
        self._auto_resize_input_box()

    def _add_attachment(self, filepath):
        if not filepath or not os.path.exists(filepath):
            return
        kind = classify_attachment(filepath)
        if kind == 'blocked':
            self.append_to_chat('Sistema', '⚠️ Este asistente no procesa imágenes ni PDF. Para ese tipo de archivos usa Microsoft Copilot.', persist=False)
            return
        if kind == 'unknown':
            self.append_to_chat('Sistema', f'⚠️ Tipo de archivo no soportado: {os.path.basename(filepath)}', persist=False)
            return
        if any(a.get('path') == filepath for a in self.pending_attachments):
            return
        self.pending_attachments.append({'id': uuid.uuid4().hex, 'path': filepath, 'name': os.path.basename(filepath), 'type': kind, 'loaded': False})

    def _remove_attachment(self, att_id):
        self.pending_attachments = [a for a in self.pending_attachments if a.get('id') != att_id]
        self._refresh_attachments_bar()
        self._auto_resize_input_box()

    def _refresh_attachments_bar(self):
        for w in self.attachments_bar.winfo_children():
            w.destroy()
        if not self.pending_attachments:
            self.attachments_bar.grid_remove()
            return
        self.attachments_bar.grid(row=0, column=0, sticky='ew', padx=8, pady=(6, 0))
        for att in self.pending_attachments:
            icon = '📊' if att['type'] == 'excel' else '📄'
            chip = tk.Frame(self.attachments_bar, bg=COLORS['chip'], bd=0)
            chip.pack(side=tk.LEFT, padx=(0, 6), pady=(0, 4))
            label_text = f"{icon} {first_nonempty_line(att['name'], 28)}"
            tk.Label(chip, text=label_text, bg=COLORS['chip'], fg=COLORS['text'], font=('Segoe UI', 9)).pack(side=tk.LEFT, padx=(8, 4), pady=4)
            tk.Button(chip, text='×', command=lambda aid=att['id']: self._remove_attachment(aid), bg=COLORS['chip'], fg=COLORS['muted'], activebackground=COLORS['danger'], activeforeground=COLORS['white'], relief=tk.FLAT, width=2, font=('Segoe UI', 9, 'bold')).pack(side=tk.LEFT, padx=(0, 4), pady=2)

    def _attachment_summary_text(self, attachments):
        if not attachments:
            return ''
        lines = ['Adjuntos enviados con este mensaje:']
        for att in attachments:
            lines.append(f"- {att['type'].upper()}: {att['name']} | Ruta local: {att['path']}")
        return '\n'.join(lines)

    def refresh_recent_list(self):
        for w in self.recent_frame.winfo_children():
            w.destroy()
        active_id = self.memory.get_active_conversation_id()
        conversations = self.memory.list_conversations(self.search_var.get() if hasattr(self, 'search_var') else '')
        for conv in conversations:
            bg = COLORS['sidebar_selected'] if conv.get('id') == active_id else COLORS['sidebar']
            row = tk.Frame(self.recent_frame, bg=bg)
            row.pack(fill=tk.X, pady=1)
            title = ('📌 ' if conv.get('pinned') else '') + first_nonempty_line(conv.get('title', 'Nueva conversación'), 31)
            label = tk.Label(row, text=title, bg=bg, fg=COLORS['text'], anchor='w', padx=8, font=('Segoe UI', 9), cursor='hand2')
            label.pack(side=tk.LEFT, fill=tk.X, expand=True)
            label.bind('<Button-1>', lambda e, cid=conv['id']: self.open_conversation(cid))
            tk.Button(row, text='×', command=lambda cid=conv['id']: self.delete_conversation(cid), bg=bg, fg=COLORS['muted_2'], activebackground=COLORS['danger'], activeforeground=COLORS['white'], relief=tk.FLAT, width=2, font=('Segoe UI', 9, 'bold')).pack(side=tk.RIGHT, padx=(4, 4), pady=2)

    def new_conversation(self):
        if self.is_busy:
            return
        self.memory.create_conversation(title='Nueva conversación')
        self.active_conversation = self.memory.get_conversation()
        self.history = []
        self.pending_attachments = []
        self.last_sent_attachments = []
        self.clear_chat_display()
        self._refresh_attachments_bar()
        self._auto_resize_input_box()
        self.append_to_chat('Sistema', '✨ Nueva conversación creada. Puedes cargar Excel, adjuntar archivos o preguntar directamente.', persist=False)
        self.lbl_chat_title.config(text='Nueva conversación')
        self.refresh_recent_list()

    def open_conversation(self, conv_id):
        if self.is_busy:
            return
        if self.memory.set_active_conversation(conv_id):
            self.active_conversation = self.memory.get_conversation(conv_id)
            self.history = self._messages_for_llm()
            self.pending_attachments = []
            self.last_sent_attachments = []
            self._refresh_attachments_bar()
            self._auto_resize_input_box()
            self.load_active_conversation_into_chat()
            self.refresh_recent_list()

    def delete_conversation(self, conv_id):
        if self.is_busy:
            return
        if messagebox.askyesno('Eliminar conversación', '¿Deseas eliminar esta conversación local?'):
            self.memory.delete_conversation(conv_id)
            self.active_conversation = self.memory.get_conversation()
            self.history = self._messages_for_llm()
            self.load_active_conversation_into_chat()
            self.refresh_recent_list()

    def _messages_for_llm(self):
        conv = self.memory.get_conversation()
        return [{'role': m.get('role'), 'content': m.get('content', '')} for m in conv.get('messages', []) if m.get('role') in ('user', 'assistant')][-MAX_CHAT_HISTORY:]

    def load_active_conversation_into_chat(self):
        self.clear_chat_display()
        conv = self.memory.get_conversation()
        self.lbl_chat_title.config(text=conv.get('title', 'Nueva conversación'))
        messages = conv.get('messages', [])
        if not messages:
            self.append_to_chat('Sistema', '📊 Asistente analítico iniciado. Puedes adjuntar Excel, Docx, CSV o Txt desde el botón ＋ del cuadro inferior.', persist=False)
            return
        role_to_sender = {'user': 'Tú', 'assistant': 'Gemma', 'system': 'Sistema', 'python': 'Consola Python', 'error': 'Error'}
        for msg in messages[-80:]:
            self.append_to_chat(role_to_sender.get(msg.get('role'), 'Sistema'), msg.get('content', ''), persist=False)

    def clear_chat_display(self):
        self.chat_display.config(state=tk.NORMAL)
        self.chat_display.delete('1.0', tk.END)
        self.chat_display.config(state=tk.DISABLED)

    def append_to_chat(self, sender, text, persist=False, role=None, extra=None):
        self.chat_display.config(state=tk.NORMAL)
        colors = {'Tú': COLORS['accent'], 'Gemma': COLORS['accent_2'], 'Gemma (Análisis)': COLORS['accent_2'], 'Sistema': '#c58af9', 'Consola Python': COLORS['warning'], 'Error': COLORS['danger']}
        self.chat_display.tag_config(sender, foreground=colors.get(sender, COLORS['text']), font=('Segoe UI', 11, 'bold'))
        self.chat_display.insert(tk.END, f'{sender}\n', sender)
        self.chat_display.insert(tk.END, f'{text}\n\n')
        self.chat_display.yview(tk.END)
        self.chat_display.config(state=tk.DISABLED)
        if persist:
            save_role = role or {'Tú': 'user', 'Gemma': 'assistant', 'Gemma (Análisis)': 'assistant', 'Sistema': 'system', 'Consola Python': 'python', 'Error': 'error'}.get(sender, 'system')
            self.memory.append_message(save_role, text, extra=extra)
            self.active_conversation = self.memory.get_conversation()
            self.lbl_chat_title.config(text=self.active_conversation.get('title', 'Nueva conversación'))
            self.refresh_recent_list()

    def toggle_thinking_panel(self):
        if bool(self.show_thinking_var.get()):
            self.thinking_panel_visible = True
            self.body.grid_columnconfigure(1, minsize=DEBUG_PANEL_WIDTH)
            self.debug_frame.grid(row=0, column=1, sticky='nsew')
        else:
            self.thinking_panel_visible = False
            self.debug_frame.grid_remove()
            self.body.grid_columnconfigure(1, minsize=0)
        self.root.update_idletasks()

    def clear_thinking_panel(self):
        self.thinking_display.config(state=tk.NORMAL)
        self.thinking_display.delete('1.0', tk.END)
        self.thinking_display.config(state=tk.DISABLED)
        self.lbl_thinking_status.config(text='Inactivo')

    def set_thinking_status(self, text):
        self.lbl_thinking_status.config(text=text)

    def handle_return(self, event):
        if not event.state & 0x0001:
            self.send_message()
            return 'break'

    def _auto_resize_input_box(self, event=None):
        if not hasattr(self, 'input_box'):
            return
        try:
            content = self.input_box.get('1.0', 'end-1c')
            font = tkfont.Font(font=self.input_box.cget('font'))
            line_space = max(font.metrics('linespace'), 20)
            width_px = max(self.input_box.winfo_width() - 36, 220)
            avg_char_px = max(font.measure('0'), 7)
            chars_per_line = max(int(width_px / avg_char_px), 28)
            visual_lines = 1
            for paragraph in content.split('\n'):
                visual_lines += max(1, (len(paragraph) // chars_per_line) + 1)
            target_lines = max(INPUT_MIN_LINES, min(INPUT_MAX_LINES, visual_lines))
            extra_attach = ATTACH_BAR_HEIGHT if self.pending_attachments else 0
            target_height = max(INPUT_MIN_HEIGHT + extra_attach, min(INPUT_MAX_HEIGHT + extra_attach, int(target_lines * line_space + 44 + extra_attach)))
            self.input_box.configure(height=target_lines)
            self.input_outer.configure(height=target_height)
            self.input_inner.configure(height=max(40, target_height - 2))
            self.chat_frame.grid_rowconfigure(1, minsize=target_height)
            self.chat_display.after_idle(lambda: self.chat_display.yview(tk.END))
        except Exception:
            pass

    def _start_live_answer(self, sender='Gemma'):
        if self.cancel_event.is_set() or self.live_answer_started:
            return
        self.live_answer_started = True
        self.chat_display.config(state=tk.NORMAL)
        self.chat_display.tag_config(sender, foreground=COLORS['accent_2'], font=('Segoe UI', 11, 'bold'))
        self.chat_display.insert(tk.END, f'{sender}\n', sender)
        self.chat_display.config(state=tk.DISABLED)
        self.chat_display.yview(tk.END)

    def _append_live_answer(self, chunk_text):
        if self.cancel_event.is_set():
            return
        self.chat_display.config(state=tk.NORMAL)
        self.chat_display.insert(tk.END, chunk_text)
        self.chat_display.yview(tk.END)
        self.chat_display.config(state=tk.DISABLED)

    def _finish_live_answer(self):
        if self.cancel_event.is_set():
            return
        self.chat_display.config(state=tk.NORMAL)
        self.chat_display.insert(tk.END, '\n\n')
        self.chat_display.config(state=tk.DISABLED)
        self.chat_display.yview(tk.END)
        self.live_answer_started = False
        self.last_chunk = ''

    def _update_live_thinking(self, text):
        if self.cancel_event.is_set():
            return
        self.thinking_display.config(state=tk.NORMAL)
        self.thinking_display.delete('1.0', tk.END)
        self.thinking_display.insert(tk.END, text)
        self.thinking_display.config(state=tk.DISABLED)
        self.thinking_display.yview(tk.END)

    def remember_current_note(self):
        text = self.input_box.get('1.0', tk.END).strip()
        if not text:
            self.append_to_chat('Sistema', "Escribe una nota en la caja de texto y presiona 'Recordar nota'.", persist=False)
            return
        self.memory.add_memory_note(text, source='manual')
        self.input_box.delete('1.0', tk.END)
        self._auto_resize_input_box()
        self.append_to_chat('Sistema', f'🧠 Recuerdo guardado localmente: {text}', persist=True, role='system')

    def load_excel(self):
        filepath = filedialog.askopenfilename(filetypes=[('Excel files', '*.xlsx *.xls')])
        if not filepath:
            return
        self.current_filepath = filepath
        self.btn_send.config(state=tk.DISABLED)
        self.is_busy = True
        self.append_to_chat('Sistema', f'📥 Cargando archivo: {os.path.basename(filepath)} ...', persist=True, role='system')
        threading.Thread(target=self._load_excel_worker, args=(filepath,), daemon=True).start()

    def _load_excel_worker(self, filepath):
        try:
            self._load_excel_from_path_blocking(filepath)
        except Exception as e:
            err_msg = str(e)
            self.root.after(0, lambda m=err_msg: self.append_to_chat('Error', f'Fallo al cargar: {m}', persist=True, role='error'))
        finally:
            self.root.after(0, self._release_ui)

    def _load_excel_from_path_blocking(self, filepath):
        self.current_filepath = filepath
        excel_file = pd.ExcelFile(filepath)
        sheets = excel_file.sheet_names
        sheet = sheets[0]
        if len(sheets) > 1:
            dialog_result = {'done': False, 'sheet': None}
            def open_sheet_dialog():
                d = SheetSelectorDialog(self.root, sheets)
                self.root.wait_window(d)
                dialog_result['sheet'] = d.selected_sheet
                dialog_result['done'] = True
            self.root.after(0, open_sheet_dialog)
            while not dialog_result['done'] and self.root.winfo_exists():
                time.sleep(0.1)
            sheet = dialog_result['sheet']
            if not sheet:
                return None
        temp_proc = ExcelProcessor(filepath, sheet)
        temp_proc.load()
        mem_data = self.memory.get_dataset(temp_proc.signature)
        specific_defs = mem_data.get('column_definitions', {}) if mem_data else {}
        final_defs = {}
        for col in temp_proc.df.columns:
            if col in specific_defs:
                final_defs[col] = specific_defs[col]
            else:
                global_meta = self.memory.get_global_column_meta(col)
                final_defs[col] = global_meta if global_meta else {'role': '', 'meaning': ''}
        processor = ExcelProcessor(filepath, sheet, column_definitions=final_defs)
        processor.load()
        dialog_result = {'done': False, 'result': None}
        def open_metadata_dialog():
            d = ColumnDefinitionDialog(self.root, processor)
            self.root.wait_window(d)
            dialog_result['result'] = d.result
            dialog_result['done'] = True
        self.root.after(0, open_metadata_dialog)
        while not dialog_result['done'] and self.root.winfo_exists():
            time.sleep(0.1)
        if dialog_result['result']:
            processor.apply_user_definitions(dialog_result['result'])
            self.memory.update_dataset(processor.signature, processor.build_memory_payload())
            self.memory.update_global_dictionary(dialog_result['result'])
            self.memory.remember_dataset_for_active_chat(processor.signature)
            msg = '✅ Diccionario guardado y actualizado globalmente.\n\n' + processor.build_llm_context()
        else:
            msg = 'Carga finalizada utilizando la autodetección de la herramienta.'
        self.processors.append(processor)
        self.dfs.append(processor.df)
        self.root.after(0, lambda m=msg: self.append_to_chat('Sistema', f'{m}\n\nArchivos cargados en esta sesión: {len(self.processors)}', persist=True, role='system'))
        return processor

    def _release_ui(self):
        self.is_busy = False
        self.btn_send.config(state=tk.NORMAL, text='Enviar', command=self.send_message, bg=COLORS['accent'], fg='#08111f')
        self.btn_attach.config(state=tk.NORMAL)

    def cancel_analysis(self):
        if not self.is_busy:
            return
        self.cancel_event.set()
        self.root.after(0, self._force_finish_answer)
        self.root.after(0, lambda: self.set_thinking_status('Cancelado'))
        self.root.after(0, self.clear_thinking_panel)
        self.append_to_chat('Sistema', '⛔ Análisis cancelado por el usuario. Puedes corregir o enviar un nuevo prompt.', persist=False)
        self._release_ui()

    def _force_finish_answer(self):
        try:
            self.chat_display.config(state=tk.NORMAL)
            self.chat_display.insert(tk.END, '\n\n')
            self.chat_display.config(state=tk.DISABLED)
            self.chat_display.yview(tk.END)
        except Exception:
            pass
        self.live_answer_started = False
        self.last_chunk = ''

    def send_message(self):
        if self.is_busy:
            return
        user_text = self.input_box.get('1.0', tk.END).strip()
        if not user_text and not self.pending_attachments:
            return
        attachments_to_send = list(self.pending_attachments)
        self.last_sent_attachments = attachments_to_send
        attach_summary = self._attachment_summary_text(attachments_to_send)
        display_text = user_text if user_text else 'Analiza los archivos adjuntos.'
        if attach_summary:
            display_text += '\n\n' + attach_summary
        self.input_box.delete('1.0', tk.END)
        self.pending_attachments = []
        self._refresh_attachments_bar()
        self._auto_resize_input_box()
        self.append_to_chat('Tú', display_text, persist=True, role='user', extra={'attachments': attachments_to_send} if attachments_to_send else None)
        self.current_run_id += 1
        run_id = self.current_run_id
        self.cancel_event.clear()
        self.is_busy = True
        self.btn_send.config(text='Cancelar', command=self.cancel_analysis, bg=COLORS['danger'], fg=COLORS['white'])
        self.btn_attach.config(state=tk.DISABLED)
        self.history = self._messages_for_llm()
        threading.Thread(target=self._orchestrate_ai_workflow, args=(run_id,), daemon=True).start()

    def _process_last_sent_attachments(self):
        if not self.last_sent_attachments:
            return ''
        context_blocks = []
        for att in self.last_sent_attachments:
            filepath = att.get('path')
            self.current_filepath = filepath
            filename = att.get('name', os.path.basename(filepath) if filepath else 'archivo')
            kind = att.get('type')
            if not filepath or not os.path.exists(filepath):
                context_blocks.append(f'Archivo no encontrado: {filename}')
                continue
            if kind == 'excel':
                self.root.after(0, lambda n=filename: self.append_to_chat('Sistema', f'📎 Procesando Excel adjunto: {n}', persist=True, role='system'))
                self._load_excel_from_path_blocking(filepath)
                code = build_reader_code(filepath)
                success, output = self._run_reader_sandbox(code)
                context_blocks.append(f'Archivo Excel procesado: {filename}\nRuta local disponible como filepath: {filepath}\n\n{output}' if success else f'Error extrayendo Excel {filename}:\n{output}')
            elif kind in ('csv', 'text', 'word'):
                self.root.after(0, lambda n=filename: self.append_to_chat('Sistema', f'📎 Extrayendo contenido del archivo: {n}', persist=False))
                code = build_reader_code(filepath)
                success, output = self._run_reader_sandbox(code)
                context_blocks.append(f'Archivo procesado: {filename}\nTipo: {kind}\nRuta local disponible como filepath: {filepath}\n\nContenido extraído por Python:\n{output}' if success else f'Error procesando {filename}:\n{output}')
            else:
                context_blocks.append(f'Archivo omitido: {filename}. Tipo no soportado para lectura automática.')
        self.last_sent_attachments = []
        return '\n\n---\n\n'.join(context_blocks) if context_blocks else ''

    def _orchestrate_ai_workflow(self, run_id):
        try:
            if run_id != self.current_run_id:
                return
            max_attempts = 6
            attempt = 0
            attachment_context = self._process_last_sent_attachments()
            if attachment_context:
                self.history.append({'role': 'user', 'content': 'Contexto de archivos adjuntos procesados por Python:\n' + attachment_context})
            ai_reply = self._stream_ollama_call(run_id=run_id)
            if self.cancel_event.is_set() or run_id != self.current_run_id:
                return
            if ai_reply.strip():
                self.memory.append_message('assistant', ai_reply)
                self.history = self._messages_for_llm()
            while attempt < max_attempts:
                if self.cancel_event.is_set() or run_id != self.current_run_id:
                    return
                code_blocks = re.findall(r'```python(.*?)```', ai_reply, re.DOTALL | re.IGNORECASE)
                if not code_blocks:
                    break
                success = True
                output = ''
                fig = None
                for code_block in code_blocks:
                    if self.cancel_event.is_set() or run_id != self.current_run_id:
                        return
                    success, output, fig = self._run_python_sandbox(code_block.strip())
                    if output:
                        self.root.after(0, lambda o=output, ok=success: self.append_to_chat('Consola Python' if ok else 'Error', o, persist=True, role='python' if ok else 'error'))
                    if fig is not None:
                        self.root.after(0, lambda f=fig: self._show_figure(f))
                    if not success:
                        break
                if success:
                    break
                attempt += 1
                feedback_prompt = 'El código anterior falló al ejecutarse en COPA AI.\n\nError o salida recibida:\n' + output + '\n\nGenera una versión corregida del código en un único bloque ```python```.'
                self.history.append({'role': 'user', 'content': feedback_prompt})
                ai_reply = self._stream_ollama_call(is_feedback=True, run_id=run_id)
                if self.cancel_event.is_set() or run_id != self.current_run_id:
                    return
                if ai_reply.strip():
                    self.memory.append_message('assistant', ai_reply)
                    self.history = self._messages_for_llm()
        except Exception as e:
            self.root.after(0, lambda m=str(e): self.append_to_chat('Error', m, persist=True, role='error'))
        finally:
            self.root.after(0, self._release_ui)

    def _run_python_sandbox(self, code_string):
        import traceback
        import webbrowser
        output_buffer = io.StringIO()
        error_buffer = io.StringIO()
        filepath = getattr(self, 'filepath', None) or getattr(self, 'current_filepath', None) or getattr(self, 'attached_filepath', None) or getattr(self, 'selected_filepath', None) or None
        if not getattr(self, 'dfs', None) and not filepath:
            return False, 'Error: No hay ningún DataFrame cargado ni ruta de archivo disponible.', None
        try:
            plt.close('all')
        except Exception:
            pass
        custom_show = {'called': False}
        original_plt_show = plt.show
        def mock_show(*args, **kwargs):
            custom_show['called'] = True
            return None
        try:
            import plotly.io as pio
            pio.renderers.default = 'browser'
        except Exception:
            pio = None
        dfs = getattr(self, 'dfs', []) or []
        df = dfs[-1] if dfs else None
        env = {'__name__': '__main__', 'dfs': dfs, 'df': df, 'pd': pd, 'plt': plt, 'os': os, 'webbrowser': webbrowser, 'filepath': filepath, 'Document': Document}
        if pio is not None:
            env['pio'] = pio
        plt.show = mock_show
        try:
            plotly_patch = '''
try:
    import plotly.io as pio
    pio.renderers.default = "browser"
except Exception:
    pass
'''
            final_code = plotly_patch + '\n' + code_string
            with contextlib.redirect_stdout(output_buffer), contextlib.redirect_stderr(error_buffer):
                exec(final_code, env)
            stdout_text = output_buffer.getvalue()
            stderr_text = error_buffer.getvalue()
            full_output = stdout_text
            if stderr_text.strip():
                full_output += '\n\nSTDERR:\n' + stderr_text
            fig = None
            try:
                fig_nums = plt.get_fignums()
                if fig_nums:
                    fig = plt.figure(fig_nums[-1])
                elif custom_show['called'] or plt.gcf().get_axes():
                    fig = plt.gcf()
            except Exception:
                fig = None
            if not full_output.strip():
                full_output = 'Código ejecutado correctamente, pero no produjo salida de consola.'
            return True, full_output, fig
        except Exception:
            stdout_text = output_buffer.getvalue()
            stderr_text = error_buffer.getvalue()
            traceback_text = traceback.format_exc()
            full_error = ''
            if stdout_text.strip():
                full_error += 'Salida antes del error:\n' + stdout_text + '\n\n'
            if stderr_text.strip():
                full_error += 'STDERR:\n' + stderr_text + '\n\n'
            full_error += 'Error ejecutando código:\n' + traceback_text
            return False, full_error, None
        finally:
            plt.show = original_plt_show

    def _build_memory_context(self):
        notes = self.memory.get_memory_notes()
        if not notes:
            return ''
        lines = ['Recuerdos globales guardados localmente por el usuario:']
        for n in notes[-12:]:
            lines.append(f"- {n.get('text', '')}")
        return '\n'.join(lines)

    def _stream_ollama_call(self, is_feedback=False, run_id=None):
        messages = [{'role': 'system', 'content': self.system_prompt_text}]
        memory_context = self._build_memory_context()
        if memory_context:
            messages.append({'role': 'system', 'content': memory_context})
        if self.processors:
            ctx_msg = 'Metadatos de los archivos cargados en esta sesión:\n'
            for i, p in enumerate(self.processors):
                ctx_msg += f'\n--- ARCHIVO dfs[{i}]: {p.filename} ---\n{p.build_llm_context()}\n'
            messages.append({'role': 'system', 'content': ctx_msg})
        messages.extend([m for m in self.history if m.get('role') in ('user', 'assistant')])
        show_think = bool(self.show_thinking_var.get())
        think_text = ''
        ans_text = ''
        if show_think and self.thinking_panel_visible:
            self.root.after(0, self.clear_thinking_panel)
        self.root.after(0, lambda: self.set_thinking_status('Analizando...' if not is_feedback else 'Concluyendo...'))
        sender = 'Gemma' if not is_feedback else 'Gemma (Análisis)'
        if not self.cancel_event.is_set() and run_id == self.current_run_id:
            self.root.after(0, lambda s=sender: self._start_live_answer(sender=s))
        try:
            stream = ollama.chat(model=DEFAULT_MODEL, messages=messages, think=show_think, stream=True, options={'temperature': 0.1})
            update_counter = 0
            for chunk in stream:
                if self.cancel_event.is_set() or run_id != self.current_run_id:
                    self.root.after(0, lambda: self.set_thinking_status('Cancelado'))
                    self.root.after(0, self.clear_thinking_panel)
                    return ''
                msg = get_chunk_message_dict(chunk)
                ch_think = msg.get('thinking', '')
                if ch_think:
                    think_text += ch_think
                    if len(think_text) > 5000:
                        think_text = think_text[-5000:]
                    if show_think and self.thinking_panel_visible and update_counter % 8 == 0:
                        self.root.after(0, lambda t=think_text: self._update_live_thinking(t))
                        self.root.after(0, lambda: self.set_thinking_status('Thinking...'))
                ch_content = msg.get('content', '')
                if ch_content and ch_content != self.last_chunk:
                    self.last_chunk = ch_content
                    ans_text += ch_content
                    self.root.after(0, lambda t=ch_content: self._append_live_answer(t))
                update_counter += 1
        except Exception as e:
            if not self.cancel_event.is_set() and run_id == self.current_run_id:
                err_ans = f'[Error: {e}]'
                ans_text += err_ans
                self.root.after(0, lambda m=err_ans: self._append_live_answer(m))
        if not self.cancel_event.is_set() and run_id == self.current_run_id:
            self.root.after(0, lambda: self.set_thinking_status('Completado'))
            self.root.after(0, self._finish_live_answer)
        return ans_text

    def _show_figure(self, fig):
        win = tk.Toplevel(self.root)
        win.title('Gráfico dinámico de Pandas')
        win.geometry('900x560')
        win.configure(bg=COLORS['bg'])
        canvas = FigureCanvasTkAgg(fig, master=win)
        canvas.draw()
        canvas.get_tk_widget().pack(expand=True, fill=tk.BOTH)

    def on_close(self):
        try:
            self.memory.save()
        except Exception:
            pass
        self.root.destroy()


if __name__ == '__main__':
    root = tk.Tk()
    app = FinancialAssistantApp(root)
    root.mainloop()
